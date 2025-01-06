'''
CV Final Project: Felicia Drysdale: U43237407
This project consists of finetuning YOLOv8 for segmentation with images I took and testing them with my test
set and NFL Snapshots. The goal was to determine if it was a touchdown based on the pixel locations of the
line and the football relative to eachother. 
'''

import json
from pathlib import Path
import cv2
import numpy as np
from ultralytics import YOLO
from docx import Document
from docx.shared import Inches
import logging

#Get rid of YOLO logging
logging.getLogger("ultralytics").setLevel(logging.WARNING)

#YOLOv8 segmentation model
model = YOLO('train/weights/best.pt')

#Define paths for folders
base_folder_paths = [
    'test_felicia',
    'test_nfl'
]

#Mappings for labels
class_names = ['football', 'line']
football_class_id = 0
line_class_id = 1

#Helper function to get the most confident detections
def get_most_confident_detections(predictions):
    football_detection = None
    line_detection = None

    for det in predictions:
        cls = int(det['class'])

        #football
        if cls == football_class_id:
            if not football_detection or det['confidence'] > football_detection['confidence']:
                football_detection = det
        
        #line
        elif cls == line_class_id:
            if not line_detection or det['confidence'] > line_detection['confidence']:
                line_detection = det

    return [football_detection, line_detection]

#Determines rightmost and leftmost pixels positions from masks
def get_extreme_pixels(mask, direction='right'):
    if direction == 'right':
        return max(np.where(mask > 0)[1])
    elif direction == 'left':
        return min(np.where(mask > 0)[1])
    else:
        raise ValueError("Direction must be 'right' or 'left'.")

#Processes a single image
def process_image(image_path, output_folder):
    img = cv2.imread(str(image_path))
    if img is None:
        return None

    #Predict with model
    predictions = model(img)
    detections = []

    #Bounding box details 
    for result in predictions[0].boxes.data.tolist():
        x1, y1, x2, y2, conf, cls = result
        detections.append({
            'x1': x1,
            'y1': y1,
            'x2': x2,
            'y2': y2,
            'confidence': conf,
            'class': int(cls),
            'name': class_names[int(cls)]
        })

    filtered_detections = get_most_confident_detections(detections)

    #Check if both are detected
    if not all(filtered_detections):
        result_text = "Either no football or no line detected. Cannot determine touchdown."
        return {'image_name': image_path.name, 'result': result_text}

    masks = predictions[0].masks.data.cpu().numpy()
    orig_shape = predictions[0].masks.orig_shape
    resized_masks = [cv2.resize(mask, (orig_shape[1], orig_shape[0])) for mask in masks]

    #Extract masks
    football_mask = resized_masks[int(filtered_detections[0]['class'])]
    line_mask = resized_masks[int(filtered_detections[1]['class'])]

    #Get furthest pixels
    football_rightmost = get_extreme_pixels(football_mask, direction='right')
    line_leftmost = get_extreme_pixels(line_mask, direction='left')

    if football_rightmost > line_leftmost:
        result_text = f"Touchdown! Football crosses the line by {football_rightmost - line_leftmost} pixels."
    else:
        result_text = f"No touchdown. Football is {line_leftmost - football_rightmost} pixels away from the line."

    #Save annotated images
    img_with_boxes = img.copy()
    img_with_masks = img.copy()

    for det in filtered_detections:
        x1, y1, x2, y2 = int(det['x1']), int(det['y1']), int(det['x2']), int(det['y2'])
        label = f"{det['name']} {det['confidence']:.2f}"
        color = (0, 255, 0) if det['class'] == football_class_id else (255, 0, 0)

        #Draw bounding boxes
        cv2.rectangle(img_with_boxes, (x1, y1), (x2, y2), color, 2)
        cv2.putText(img_with_boxes, label, (x1, y1 - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 2)

        #Overlay masks for display
        mask = resized_masks[int(det['class'])]
        mask = (mask > 0).astype(np.uint8) * 255
        overlay = np.zeros_like(img_with_masks)
        overlay[mask == 255] = (0, 255, 0) if det['class'] == football_class_id else (0, 0, 255)
        img_with_masks = cv2.addWeighted(img_with_masks, 1, overlay, 0.5, 0)

    #Paths to save images
    bounding_box_path = output_folder / f"{image_path.stem}_boxes.png"
    mask_path = output_folder / f"{image_path.stem}_masks.png"

    cv2.imwrite(str(bounding_box_path), img_with_boxes)
    cv2.imwrite(str(mask_path), img_with_masks)

    return {
        'image_name': image_path.name,
        'result': result_text,
        'bounding_box_image': str(bounding_box_path),
        'mask_image': str(mask_path),
        'detections': filtered_detections
    }

#Function to create a report
def create_report(folder_path, output_file_path, has_subfolders=True):
    document = Document()
    document.add_heading(f"Results For {folder_path.name}", level=1)

    output_folder = folder_path / "OUTPUTS_RESULTS"
    output_folder.mkdir(parents=True, exist_ok=True)

    image_paths = []
    if has_subfolders:
        for subfolder in folder_path.iterdir():
            if subfolder.is_dir():

                #Iterate over all subfolders and get images
                image_paths.extend(subfolder.rglob("*.png"))
                image_paths.extend(subfolder.rglob("*.jpeg"))  
    else:
        image_paths.extend(folder_path.glob("*.png"))
        image_paths.extend(folder_path.glob("*.jpeg")) 

    #Process each image in the folder
    for image_path in image_paths:
        result = process_image(image_path, output_folder)
        if not result:
            continue

        document.add_heading(result['image_name'], level=2)
        document.add_paragraph(result['result'])

        #Embed bounding box image
        if 'bounding_box_image' in result:
            document.add_paragraph("Bounding Box Image:")
            document.add_picture(result['bounding_box_image'], width=Inches(4))

        #Embed mask image
        if 'mask_image' in result:
            document.add_paragraph("Mask Image:")
            document.add_picture(result['mask_image'], width=Inches(4))

        # Add JSON data
        if 'detections' in result:
            document.add_paragraph("JSON Data:")
            json_data = json.dumps(result['detections'], indent=4)
            document.add_paragraph(json_data)

    document.save(output_file_path)

#Process each folder
for folder_path in base_folder_paths:
    folder_path = Path(folder_path)
    has_subfolders = any(subfolder.is_dir() for subfolder in folder_path.iterdir())
    output_file_path = folder_path.parent / f"{folder_path.name}_results.docx"
    create_report(folder_path, output_file_path, has_subfolders=has_subfolders)

#All my alerting print statements
print("Reports generated successfully!")
print("Please Note: If nothing was able to be detected, it is mentioned in the report.")
print("Please Note: images with word 'bad' in them are NOT correct but still wanted to show.")
print("All output images can also be found in output folders under tests.")
